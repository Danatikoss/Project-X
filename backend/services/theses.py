"""
Theses (Talking Points) Generation Service.

Flow:
  1. create_session()   — create a ThesesSession from an assembly (snapshot slides)
  2. analyze_session()  — GPT-4o reads the snapshot and returns 2-3 clarifying questions
  3. generate_session() — GPT-4o vision analyzes each slide thumbnail + metadata and
                          produces 3-5 talking points per slide in KK / RU / EN

Style: официально-деловой, простые слова, без тяжёлых канцеляризмов.
"""
import base64
import io
import json
import logging
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from openai import AsyncOpenAI
from sqlalchemy.orm import Session

from config import settings
from models.assembly import AssembledPresentation
from models.slide import SlideLibraryEntry
from models.theses import ThesesSession

logger = logging.getLogger(__name__)

# ── Prompts ───────────────────────────────────────────────────────────────────

ANALYZE_SYSTEM = """Ты — помощник по подготовке деловых выступлений.
Тебе дан список слайдов презентации (названия и краткие описания).
Задача: задать 2-3 уточняющих вопроса, которые помогут составить более точные тезисы.

Вопросы должны касаться:
- аудитории (кто слушает?)
- цели выступления
- акцентов (что особо подчеркнуть?)

Верни строго JSON:
{
  "questions": [
    {"id": "audience", "text": "Кто будет слушать выступление?"},
    {"id": "goal",     "text": "Какова главная цель вашего выступления?"},
    {"id": "emphasis", "text": "Есть ли темы или слайды, которые нужно выделить особо?"}
  ]
}

Ответь ТОЛЬКО JSON без пояснений."""

GENERATE_SYSTEM = """Ты — профессиональный бизнес-спичрайтер. Составляешь тезисы для выступления к каждому слайду презентации.

Требования к тезисам:
- Стиль: официально-деловой, но живой — без тяжёлых канцеляризмов
- Простые, понятные слова
- 3-5 тезисов на слайд (каждый — одно законченное предложение)
- Тезисы — то, что спикер ГОВОРИТ вслух, а не описание слайда
- Переведи тезисы на казахский (kk), русский (ru) и английский (en)

{context_block}

Верни строго JSON:
{{
  "<slide_id>": {{
    "ru": ["Тезис 1", "Тезис 2", "Тезис 3"],
    "kk": ["Тезис 1 на казахском", ...],
    "en": ["Thesis 1 in English", ...]
  }},
  ...
}}

Используй реальные slide_id (числа в виде строк) из запроса. Ответь ТОЛЬКО JSON."""


def _context_block(context: dict) -> str:
    if not context:
        return ""
    labels = {"audience": "Аудитория", "goal": "Цель", "emphasis": "Акценты"}
    lines = ["Контекст выступления:"]
    for k, v in context.items():
        if v and str(v).strip():
            lines.append(f"- {labels.get(k, k)}: {v}")
    return "\n".join(lines) if len(lines) > 1 else ""


def _get_client() -> AsyncOpenAI:
    kwargs: dict = {"api_key": settings.openai_api_key}
    if settings.openai_base_url:
        kwargs["base_url"] = settings.openai_base_url
    return AsyncOpenAI(**kwargs)


def _resize_thumbnail(img_bytes: bytes, max_width: int = 768) -> bytes:
    from PIL import Image
    img = Image.open(io.BytesIO(img_bytes))
    if img.width > max_width:
        ratio = max_width / img.width
        img = img.resize((max_width, int(img.height * ratio)), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, "JPEG", quality=80)
    return buf.getvalue()


def _thumb_b64(thumbnail_path: str) -> Optional[str]:
    if not thumbnail_path:
        return None
    path = Path(settings.thumbnail_dir) / thumbnail_path
    if not path.exists():
        return None
    try:
        small = _resize_thumbnail(path.read_bytes())
        return base64.b64encode(small).decode()
    except Exception:
        return None


# ── Public API ────────────────────────────────────────────────────────────────

def create_session(db: Session, owner_id: int, assembly_id: int) -> ThesesSession:
    """
    Create a new ThesesSession from an existing assembly.
    Snapshots the current slide list so the session is self-contained.
    """
    assembly = db.query(AssembledPresentation).get(assembly_id)
    if not assembly:
        raise ValueError(f"Assembly {assembly_id} not found")
    if assembly.owner_id != owner_id:
        raise PermissionError("No access to this assembly")

    slide_ids: list[int] = json.loads(assembly.slide_ids_json or "[]")
    slides = [s for sid in slide_ids if (s := db.query(SlideLibraryEntry).get(sid))]

    snapshot = [
        {
            "id": s.id,
            "title": s.title,
            "summary": s.summary,
            "thumbnail_path": s.thumbnail_path,
            "layout_type": s.layout_type,
            "tags": json.loads(s.tags_json or "[]"),
        }
        for s in slides
    ]

    session = ThesesSession(
        owner_id=owner_id,
        title=assembly.title,
        assembly_id=assembly_id,
        slide_snapshot_json=json.dumps(snapshot, ensure_ascii=False),
    )
    db.add(session)
    db.commit()
    db.refresh(session)
    return session


def create_session_from_upload(
    db: Session,
    owner_id: int,
    title: str,
    file_bytes: bytes,
    file_ext: str,
) -> ThesesSession:
    """
    Create a ThesesSession by extracting slides from an uploaded PPTX or PDF file.
    Thumbnails are saved to thumbnail_dir using the session id as a namespace.
    """
    from services.thumbnail import extract_pptx_slides, extract_pdf_slides, save_thumbnail

    with tempfile.NamedTemporaryFile(suffix=f".{file_ext}", delete=False) as f:
        f.write(file_bytes)
        tmp_path = f.name

    try:
        if file_ext == "pptx":
            slides_data = extract_pptx_slides(tmp_path)
        else:
            slides_data = extract_pdf_slides(tmp_path)
    finally:
        os.unlink(tmp_path)

    # Create session first to get a stable ID for thumbnail namespacing
    session = ThesesSession(
        owner_id=owner_id,
        title=title,
        assembly_id=None,
        slide_snapshot_json="[]",
    )
    db.add(session)
    db.commit()
    db.refresh(session)

    # Use a large offset so thumbnail dirs don't collide with real source IDs
    thumb_source_id = 10_000_000 + session.id
    snapshot = []
    for i, slide in enumerate(slides_data):
        thumb_path = save_thumbnail(slide.thumbnail_bytes, thumb_source_id, i, settings.thumbnail_dir)
        snapshot.append({
            "id": i + 1,
            "title": slide.pptx_title or f"Слайд {i + 1}",
            "summary": (slide.text or "")[:200],
            "thumbnail_path": thumb_path,
            "layout_type": None,
            "tags": [],
        })

    session.slide_snapshot_json = json.dumps(snapshot, ensure_ascii=False)
    db.commit()
    return session


def create_session_from_docx(
    db: Session,
    owner_id: int,
    title: str,
    file_bytes: bytes,
) -> ThesesSession:
    """
    Create a ThesesSession from a DOCX file containing pre-written theses.
    Headings become slide titles; paragraphs under each heading become bullet theses.
    The session is created with theses_json already populated (ru only) so the user
    lands directly on the result page and can translate/regenerate as needed.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))

    sections: list[dict] = []
    current: dict | None = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if current:
                sections.append(current)
            current = {"title": text, "bullets": []}
        else:
            if current is None:
                current = {"title": title, "bullets": []}
            current["bullets"].append(text)

    if current:
        sections.append(current)

    if not sections:
        all_bullets = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        sections = [{"title": title, "bullets": all_bullets}]

    snapshot = []
    theses_dict: dict = {}

    for i, section in enumerate(sections):
        slide_id = i + 1
        snapshot.append({
            "id": slide_id,
            "title": section["title"],
            "summary": " ".join(section["bullets"][:2]),
            "thumbnail_path": None,
            "layout_type": "content",
            "tags": [],
        })
        theses_dict[str(slide_id)] = {
            "ru": section["bullets"],
            "kk": [],
            "en": [],
        }

    session = ThesesSession(
        owner_id=owner_id,
        title=title,
        assembly_id=None,
        slide_snapshot_json=json.dumps(snapshot, ensure_ascii=False),
        theses_json=json.dumps(theses_dict, ensure_ascii=False),
    )
    db.add(session)
    db.commit()
    db.refresh(session)
    return session


async def analyze_session(db: Session, session_id: int) -> dict:
    """
    Analyze the presentation slides and return 2-3 clarifying questions.
    Returns {"questions": [{"id": "...", "text": "..."}, ...]}.
    """
    session = db.query(ThesesSession).get(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    snapshot: list[dict] = json.loads(session.slide_snapshot_json or "[]")

    lines = [f"Презентация: «{session.title}»", f"Всего слайдов: {len(snapshot)}", ""]
    for i, s in enumerate(snapshot, 1):
        title = s.get("title") or f"Слайд {i}"
        summary = s.get("summary") or ""
        lines.append(f"{i}. {title}" + (f" — {summary}" if summary else ""))

    client = _get_client()
    try:
        resp = await client.chat.completions.create(
            model=settings.assembly_model,
            messages=[
                {"role": "system", "content": ANALYZE_SYSTEM},
                {"role": "user", "content": "\n".join(lines)},
            ],
            temperature=0.3,
            response_format={"type": "json_object"},
        )
        return json.loads(resp.choices[0].message.content or "{}")
    except Exception as e:
        logger.warning(f"analyze_session failed: {e}")
        return {
            "questions": [
                {"id": "audience", "text": "Кто будет слушать выступление?"},
                {"id": "goal",     "text": "Какова главная цель вашего выступления?"},
                {"id": "emphasis", "text": "Есть ли темы, которые нужно выделить особо?"},
            ]
        }


async def generate_session(
    db: Session,
    session_id: int,
    context: Optional[dict] = None,
) -> dict:
    """
    Generate talking-point theses per slide in KK/RU/EN.
    Saves result to ThesesSession. Returns theses dict.
    """
    session = db.query(ThesesSession).get(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    context = context or {}
    snapshot: list[dict] = json.loads(session.slide_snapshot_json or "[]")
    if not snapshot:
        raise ValueError("Нет слайдов для генерации тезисов")

    system_prompt = GENERATE_SYSTEM.format(context_block=_context_block(context))

    # Build user message
    user_content: list = []
    info_lines = [f"Презентация: «{session.title}»\n"]
    for s in snapshot:
        tags = s.get("tags") or []
        info_lines.append(
            f"slide_id={s['id']}: «{s.get('title') or '(без названия)'}» | "
            f"Тип: {s.get('layout_type') or '?'} | "
            f"Теги: {', '.join(tags) if tags else '—'} | "
            f"{(s.get('summary') or '')[:120]}"
        )
    user_content.append({"type": "text", "text": "\n".join(info_lines)})

    # Add thumbnails for visual context (up to 12 slides)
    for s in snapshot[:12]:
        b64 = _thumb_b64(s.get("thumbnail_path") or "")
        if b64:
            user_content.append({"type": "text", "text": f"Слайд {s['id']} — «{s.get('title') or '?'}»:"})
            user_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "low"},
            })

    client = _get_client()
    resp = await client.chat.completions.create(
        model=settings.assembly_model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        temperature=0.4,
        response_format={"type": "json_object"},
        max_tokens=4000,
    )
    raw = resp.choices[0].message.content or "{}"
    try:
        theses = json.loads(raw)
    except json.JSONDecodeError as e:
        logger.error(f"generate_session: invalid JSON from model: {e}\nRaw: {raw[:500]}")
        raise ValueError("Модель вернула некорректный JSON. Попробуйте ещё раз.")

    session.theses_json = json.dumps(theses, ensure_ascii=False)
    session.context_json = json.dumps(context, ensure_ascii=False)
    session.updated_at = datetime.now(timezone.utc)
    db.commit()

    return theses


def list_sessions(db: Session, owner_id: int) -> list[dict]:
    """List all theses sessions for a user, newest first."""
    rows = (
        db.query(ThesesSession)
        .filter(ThesesSession.owner_id == owner_id)
        .order_by(ThesesSession.updated_at.desc())
        .limit(50)
        .all()
    )
    result = []
    for s in rows:
        snapshot = json.loads(s.slide_snapshot_json or "[]")
        theses = json.loads(s.theses_json or "{}")
        result.append({
            "id": s.id,
            "title": s.title,
            "assembly_id": s.assembly_id,
            "slide_count": len(snapshot),
            "has_theses": bool(theses),
            "created_at": s.created_at.isoformat() if s.created_at else None,
            "updated_at": s.updated_at.isoformat() if s.updated_at else None,
            # first 3 thumbnails for preview
            "thumbnail_paths": [
                sn["thumbnail_path"]
                for sn in snapshot[:3]
                if sn.get("thumbnail_path")
            ],
        })
    return result


def get_session(db: Session, session_id: int) -> Optional[dict]:
    """Return full session data including slides snapshot and theses."""
    s = db.query(ThesesSession).get(session_id)
    if not s:
        return None
    return {
        "id": s.id,
        "title": s.title,
        "assembly_id": s.assembly_id,
        "slides": json.loads(s.slide_snapshot_json or "[]"),
        "theses": json.loads(s.theses_json or "{}"),
        "context": json.loads(s.context_json or "{}"),
        "created_at": s.created_at.isoformat() if s.created_at else None,
        "updated_at": s.updated_at.isoformat() if s.updated_at else None,
    }
